"""
Compose exporter — walks section source files and produces bundled output.

Supports:
  - markdown: concatenates all .md files from sections into a single document
  - zip: creates a zip archive of all section source files
  - docx: generates a formatted Word document with section headings and page breaks
  - pdf: generates PDF via DOCX-to-PDF conversion (requires docx2pdf or LibreOffice)
"""

import io
import logging
import re
import tempfile
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Optional

from .models import project_dir, get_project, get_sections, _sanitize_folder_name

logger = logging.getLogger(__name__)


def export_markdown(project_id: str) -> Optional[str]:
    """Walk sections in order, concatenate .md files, return combined text."""
    sections = get_sections(project_id)
    if not sections:
        return None

    # Sort by order
    sections.sort(key=lambda s: (s.order, s.name))

    pdir = project_dir(project_id)
    parts = []

    for section in sections:
        section_dir = pdir / "sections" / _sanitize_folder_name(section.name) / "content"
        if not section_dir.is_dir():
            continue
        # Collect .md files in this section
        md_files = sorted(section_dir.glob("*.md"))
        if md_files:
            # Add section heading
            depth = "#" if not section.parent_id else "##"
            parts.append(f"{depth} {section.name}\n")
            for md_file in md_files:
                try:
                    content = md_file.read_text(encoding="utf-8").strip()
                    if content:
                        parts.append(content)
                except Exception:
                    logger.warning("Failed to read %s", md_file)
            parts.append("")  # blank line separator

    if not parts:
        return "# (No content yet)\n"

    return "\n\n".join(parts)


def export_zip(project_id: str) -> Optional[bytes]:
    """Create a zip archive of all section source files. Returns bytes."""
    sections = get_sections(project_id)
    if not sections:
        return None

    sections.sort(key=lambda s: (s.order, s.name))
    pdir = project_dir(project_id)

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        file_count = 0
        for section in sections:
            folder_name = _sanitize_folder_name(section.name)
            section_dir = pdir / "sections" / folder_name / "content"
            if not section_dir.is_dir():
                continue
            for fp in sorted(section_dir.iterdir()):
                if fp.is_file():
                    try:
                        arcname = f"{folder_name}/{fp.name}"
                        zf.write(fp, arcname)
                        file_count += 1
                    except Exception:
                        logger.warning("Failed to add %s to zip", fp)

        # Also include the combined markdown as master.md
        combined = export_markdown(project_id)
        if combined:
            zf.writestr("master.md", combined.encode("utf-8"))
            file_count += 1

    if file_count == 0:
        return None

    return buf.getvalue()


# ---------------------------------------------------------------------------
# DOCX Export
# ---------------------------------------------------------------------------

def _collect_section_content(project_id: str, section) -> str:
    """Read all .md content files for a section, return combined text."""
    pdir = project_dir(project_id)
    section_dir = pdir / "sections" / _sanitize_folder_name(section.name) / "content"
    if not section_dir.is_dir():
        return ""
    parts = []
    for md_file in sorted(section_dir.glob("*.md")):
        try:
            content = md_file.read_text(encoding="utf-8").strip()
            if content:
                parts.append(content)
        except Exception:
            logger.warning("Failed to read %s", md_file)
    return "\n\n".join(parts)


def _add_markdown_to_doc(doc, text):
    """Convert basic markdown formatting to python-docx paragraphs.

    Handles: headings (##, ###), bold, italic, inline code, bullet lists,
    numbered lists, and code blocks. Falls back to plain text for anything else.
    """
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    lines = text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]

        # Code block
        if line.strip().startswith('```'):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1  # skip closing ```
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run('\n'.join(code_lines))
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
            continue

        # Headings
        if line.startswith('#### '):
            doc.add_heading(line[5:].strip(), level=4)
            i += 1
            continue
        if line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue
        if line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue

        # Bullet list
        if re.match(r'^[\-\*]\s+', line):
            text_content = re.sub(r'^[\-\*]\s+', '', line)
            p = doc.add_paragraph(style='List Bullet')
            _add_inline_formatting(p, text_content)
            i += 1
            continue

        # Numbered list
        m = re.match(r'^(\d+)\.\s+', line)
        if m:
            text_content = line[m.end():]
            p = doc.add_paragraph(style='List Number')
            _add_inline_formatting(p, text_content)
            i += 1
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        _add_inline_formatting(p, line)
        i += 1


def _add_inline_formatting(paragraph, text):
    """Parse inline markdown (bold, italic, code) and add runs to paragraph."""
    from docx.shared import Pt, RGBColor

    # Pattern: **bold**, *italic*, `code`
    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)')
    last_end = 0
    for match in pattern.finditer(text):
        # Add text before the match
        if match.start() > last_end:
            paragraph.add_run(text[last_end:match.start()])

        if match.group(2):  # **bold**
            run = paragraph.add_run(match.group(2))
            run.bold = True
        elif match.group(3):  # *italic*
            run = paragraph.add_run(match.group(3))
            run.italic = True
        elif match.group(4):  # `code`
            run = paragraph.add_run(match.group(4))
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

        last_end = match.end()

    # Add remaining text
    if last_end < len(text):
        paragraph.add_run(text[last_end:])


def export_docx(project_id: str) -> Optional[bytes]:
    """Generate a formatted Word document from the composition.

    Returns DOCX file content as bytes, or None if no content.
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
    except ImportError:
        logger.error("python-docx not installed. Install with: pip install python-docx")
        return None

    project = get_project(project_id)
    if not project:
        return None

    sections = get_sections(project_id)
    sections.sort(key=lambda s: (s.order, s.name))

    doc = Document()

    # ── Page setup ──
    for section_obj in doc.sections:
        section_obj.top_margin = Inches(1)
        section_obj.bottom_margin = Inches(1)
        section_obj.left_margin = Inches(1)
        section_obj.right_margin = Inches(1)

    # ── Default font ──
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.15

    # ── Title page ──
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(120)
    run = title.add_run(project.name)
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(datetime.now().strftime('%B %d, %Y'))
    date_run.font.size = Pt(12)
    date_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # Page break after title
    doc.add_page_break()

    # ── Sections ──
    if not sections:
        p = doc.add_paragraph()
        run = p.add_run('[No content yet]')
        run.italic = True
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    else:
        for idx, sec in enumerate(sections):
            # Section heading
            heading = doc.add_heading(sec.name, level=1)

            # Summary as italic intro
            if sec.summary:
                p = doc.add_paragraph()
                run = p.add_run(sec.summary)
                run.italic = True
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                run.font.size = Pt(10)

            # Section content
            content = _collect_section_content(project_id, sec)
            if content:
                _add_markdown_to_doc(doc, content)
            else:
                p = doc.add_paragraph()
                run = p.add_run('[No content yet]')
                run.italic = True
                run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

            # Page break after each section except the last
            if idx < len(sections) - 1:
                doc.add_page_break()

    # ── Page numbers (footer) ──
    try:
        for section_obj in doc.sections:
            footer = section_obj.footer
            footer.is_linked_to_previous = False
            p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Add page number field
            run = p.add_run()
            fld_char1 = run._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
            run._element.append(fld_char1)
            run2 = p.add_run()
            instr = run2._element.makeelement(qn('w:instrText'), {})
            instr.text = ' PAGE '
            run2._element.append(instr)
            run3 = p.add_run()
            fld_char2 = run3._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
            run3._element.append(fld_char2)
    except Exception:
        logger.debug("Failed to add page numbers, continuing without")

    # Save to bytes
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_pdf(project_id: str) -> Optional[bytes]:
    """Generate PDF by converting DOCX output.

    Tries docx2pdf first, falls back to error message.
    Returns PDF bytes or None.
    """
    docx_bytes = export_docx(project_id)
    if not docx_bytes:
        return None

    # Write DOCX to temp file
    tmp_docx = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    tmp_pdf_path = tmp_docx.name.replace('.docx', '.pdf')
    try:
        tmp_docx.write(docx_bytes)
        tmp_docx.close()

        # Try docx2pdf
        try:
            from docx2pdf import convert
            convert(tmp_docx.name, tmp_pdf_path)
            pdf_path = Path(tmp_pdf_path)
            if pdf_path.exists():
                return pdf_path.read_bytes()
        except ImportError:
            pass
        except Exception as e:
            logger.warning("docx2pdf conversion failed: %s", e)

        # Try subprocess with LibreOffice
        try:
            import subprocess
            result = subprocess.run(
                ['soffice', '--headless', '--convert-to', 'pdf',
                 '--outdir', str(Path(tmp_docx.name).parent), tmp_docx.name],
                capture_output=True, timeout=60,
            )
            pdf_path = Path(tmp_pdf_path)
            if pdf_path.exists():
                return pdf_path.read_bytes()
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass
        except Exception as e:
            logger.warning("LibreOffice conversion failed: %s", e)

        # No converter available — return None with logged warning
        logger.warning("PDF export unavailable: install docx2pdf or LibreOffice")
        return None

    finally:
        # Clean up temp files
        try:
            Path(tmp_docx.name).unlink(missing_ok=True)
            Path(tmp_pdf_path).unlink(missing_ok=True)
        except Exception:
            pass
